import logging
import os

import fitz
from docx import Document as DocxReader
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import (
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    EMBEDDINGS_MODEL_NAME,
    VECTOR_STORE_PATH,
)
from app.utils.text_processing import preprocess_text

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(message)s")


def extract_text_from_pdf(path: str) -> str:
    text = ""
    try:
        with fitz.open(path) as doc:
            for page in doc:
                text += page.get_text("text")
        text = text.replace("\n", " ").replace("  ", " ").strip()
        return text
    except Exception as e:
        logging.error(f"Error extracting text from {path}: {e}")
        return ""


def extract_text_from_docx(path: str) -> str:
    try:
        doc = DocxReader(path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Error extracting text from {path}: {e}")
        return ""


def load_documents_from_directory(root_dir: str) -> list[Document]:
    documents = []
    supported_formats = (".pdf", ".docx")

    logging.info(f"Scanning directory: {root_dir}")

    for subdir, _, files in os.walk(root_dir):
        for file in files:
            if not file.endswith(supported_formats):
                continue

            file_path = os.path.join(subdir, file)
            logging.info(f"Processing: {file_path}")

            if file.endswith(".pdf"):
                text = extract_text_from_pdf(file_path)
            elif file.endswith(".docx"):
                text = extract_text_from_docx(file_path)
            else:
                continue

            if text.strip():
                documents.append(
                    Document(page_content=text, metadata={"source": file_path})
                )
                logging.info(f"✔ Added: {file} ({len(text)} chars)")
            else:
                logging.warning(f"⚠️  Skipped (empty): {file}")

    return documents


def build_vectorstore(
    data_root: str = "data/", db_path: str = VECTOR_STORE_PATH
) -> FAISS:
    # Validate data directory exists
    if not os.path.exists(data_root):
        raise FileNotFoundError(f"Data directory not found: {data_root}")

    logging.info("\n" + "=" * 60)
    logging.info("STARTING DOCUMENT INGESTION")
    logging.info("=" * 60)

    # Load documents
    logging.info("\n📂 Loading documents from folders...")
    docs = load_documents_from_directory(data_root)

    if not docs:
        raise ValueError(f"No documents found in {data_root}")

    logging.info(f"\n✔ Loaded {len(docs)} documents.")

    # 1. Split into chunks
    logging.info(
        f"\nSplitting documents (chunk_size={CHUNK_SIZE}, overlap={CHUNK_OVERLAP})..."
    )
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        is_separator_regex=False,
    )
    chunks = splitter.split_documents(docs)
    logging.info(f"✔ Split into {len(chunks)} chunks.")

    # 2. Preprocess chunks
    logging.info("\nPreprocessing text chunks...")
    chunks = [
        Document(
            page_content=preprocess_text(chunk.page_content),
            metadata=chunk.metadata,
        )
        for chunk in chunks
    ]

    # 2. Create embeddings
    logging.info(f"\nCreating embeddings using: {EMBEDDINGS_MODEL_NAME}")
    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDINGS_MODEL_NAME,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    # Build vectorstore
    logging.info("\n\nBuilding FAISS vectorstore...")
    vectorstore = FAISS.from_documents(chunks, embeddings)

    # Save vectorstore
    os.makedirs(os.path.dirname(db_path), exist_ok=True)
    vectorstore.save_local(db_path)

    logging.info("\n" + "=" * 60)
    logging.info(f"✔ COMPLETE! Vectorstore saved to: {db_path}")
    logging.info("=" * 60)

    return vectorstore


if __name__ == "__main__":
    build_vectorstore(data_root="data")
